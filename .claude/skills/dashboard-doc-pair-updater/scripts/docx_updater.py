#!/usr/bin/env python3
"""
Dashboard Doc-Pair Updater Helper

Reads/writes .docx files and tracks changes for dashboard documentation sync.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from datetime import datetime


def read_docx(file_path: str) -> dict:
    """
    Read a .docx file and extract all content with structure.

    Returns: {
        'filename': str,
        'paragraphs': [{'text': str, 'style': str, 'level': int}, ...],
        'tables': [{'rows': [[cell_text, ...], ...], 'header': [str, ...]}, ...],
        'full_text': str
    }
    """
    doc = Document(file_path)

    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append({
            'text': para.text,
            'style': para.style.name,
            'level': para.paragraph_format.outline_level if hasattr(para.paragraph_format, 'outline_level') else 0
        })

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cell_texts = [cell.text for cell in row.cells]
            rows.append(cell_texts)
        tables.append({'rows': rows})

    full_text = '\n'.join([p['text'] for p in paragraphs])

    return {
        'filename': Path(file_path).name,
        'file_path': file_path,
        'paragraphs': paragraphs,
        'tables': tables,
        'full_text': full_text,
        'num_paragraphs': len(paragraphs),
        'num_tables': len(tables)
    }


def write_docx(file_path: str, content: dict, original_doc_path: str) -> bool:
    """
    Write updated content to a .docx file, preserving original formatting.

    Args:
        file_path: Path to save the new .docx
        content: Updated content (paragraphs with text and style)
        original_doc_path: Path to original document (for copying formatting)

    Returns: Success boolean
    """
    try:
        # Load original document to preserve formatting
        original_doc = Document(original_doc_path)

        # Clear existing paragraphs (keep first para for style reference)
        while len(original_doc.paragraphs) > 0:
            p = original_doc.paragraphs[0]._element
            p.getparent().remove(p)

        # Add updated paragraphs
        for para_data in content.get('paragraphs', []):
            new_para = original_doc.add_paragraph(para_data['text'])
            if 'style' in para_data:
                try:
                    new_para.style = para_data['style']
                except:
                    pass  # Keep default if style doesn't exist

        # Save
        original_doc.save(file_path)
        return True
    except Exception as e:
        print(f"Error writing to {file_path}: {e}")
        return False


def find_sections(doc_content: dict, keywords: list) -> list:
    """
    Find sections/paragraphs in document that match keywords.

    Args:
        doc_content: Document dict from read_docx
        keywords: List of keywords to search for (case-insensitive)

    Returns: List of matching paragraphs with their indices
    """
    matches = []
    for i, para in enumerate(doc_content['paragraphs']):
        text_lower = para['text'].lower()
        for keyword in keywords:
            if keyword.lower() in text_lower:
                matches.append({
                    'index': i,
                    'text': para['text'],
                    'style': para['style'],
                    'keyword_matched': keyword
                })
                break  # Only match once per paragraph

    return matches


def generate_change_summary(original_path: str, updated_path: str) -> dict:
    """
    Compare original and updated documents to generate a change summary.

    Returns: {
        'sections_updated': [section names],
        'paragraphs_modified': int,
        'new_content': [text snippets of new/changed content],
        'timestamp': ISO timestamp
    }
    """
    original = read_docx(original_path)
    updated = read_docx(updated_path)

    # Simple comparison: count changed paragraphs
    changed_paragraphs = 0
    for i in range(min(len(original['paragraphs']), len(updated['paragraphs']))):
        if original['paragraphs'][i]['text'] != updated['paragraphs'][i]['text']:
            changed_paragraphs += 1

    return {
        'file': Path(updated_path).name,
        'paragraphs_modified': changed_paragraphs,
        'timestamp': datetime.now().isoformat()
    }


def list_document_structure(doc_content: dict) -> str:
    """
    Print a readable outline of the document structure (headings and sections).
    Useful for understanding what sections exist to update.
    """
    outline = []
    for i, para in enumerate(doc_content['paragraphs']):
        if 'Heading' in para['style']:
            level = para['level']
            indent = '  ' * level
            outline.append(f"{indent}[{i}] {para['text']}")

    return '\n'.join(outline) if outline else "No headings found in document"


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_updater.py <read|structure> <file_path>")
        sys.exit(1)

    command = sys.argv[1]
    file_path = sys.argv[2] if len(sys.argv) > 2 else None

    if command == "read" and file_path:
        content = read_docx(file_path)
        print(json.dumps({
            'filename': content['filename'],
            'num_paragraphs': content['num_paragraphs'],
            'num_tables': content['num_tables'],
            'first_500_chars': content['full_text'][:500]
        }, indent=2))

    elif command == "structure" and file_path:
        content = read_docx(file_path)
        print(f"Document Structure: {content['filename']}\n")
        print(list_document_structure(content))

    else:
        print("Invalid command or missing file path")
        sys.exit(1)
