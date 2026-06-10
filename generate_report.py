from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────
BLACK   = RGBColor(0x11, 0x11, 0x11)
GOLD    = RGBColor(0xE8, 0xC5, 0x47)
MID     = RGBColor(0x44, 0x44, 0x44)
LIGHT   = RGBColor(0x88, 0x88, 0x88)
RED     = RGBColor(0xC0, 0x39, 0x2B)
GREEN   = RGBColor(0x1A, 0x7A, 0x3C)
AMBER   = RGBColor(0x99, 0x6B, 0x00)
RULE    = RGBColor(0xCC, 0xCC, 0xCC)
TBL_HDR = RGBColor(0x1A, 0x1A, 0x1A)
TBL_ALT = RGBColor(0xF7, 0xF7, 0xF7)

# ── Helper: paragraph shading ─────────────────────────────────────
def shade_cell(cell, hex_str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'),   kwargs[side].get('val',   'single'))
            tag.set(qn('w:sz'),    kwargs[side].get('sz',    '4'))
            tag.set(qn('w:space'), kwargs[side].get('space', '0'))
            tag.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11, color=None, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.color.rgb = color or BLACK
    return run

def add_heading(text, level=1, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if level == 1:
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size  = Pt(10)
        run.font.color.rgb = LIGHT
        run.font.all_caps  = True
        # rule below
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size  = Pt(13)
        run.font.color.rgb = BLACK
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size  = Pt(11)
        run.font.color.rgb = MID
    return p

def add_body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = MID
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.6)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = BLACK
        r = p.add_run(text)
    else:
        r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = MID
    return p

def add_callout(text, label=None, color_hex='F7F7F0', left_color='111111'):
    """A shaded callout paragraph."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, color_hex)
    cell.width = Inches(5.8)
    # Left border accent
    set_cell_border(cell,
        left  ={'val':'single','sz':'16','color':left_color},
        top   ={'val':'none'},
        right ={'val':'none'},
        bottom={'val':'none'},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    if label:
        rb = p.add_run(label + '  ')
        rb.bold = True
        rb.font.size = Pt(10)
        rb.font.color.rgb = BLACK
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = MID
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, '111111')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(h.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row in enumerate(rows):
        drow = tbl.rows[ri + 1]
        fill = 'FFFFFF' if ri % 2 == 0 else 'F7F7F7'
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            if isinstance(val, tuple):
                text, bold, color = val
                r = p.add_run(text)
                r.bold = bold
                r.font.size = Pt(9.5)
                r.font.color.rgb = color or MID
            else:
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                r.font.color.rgb = MID
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

# ════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run('ONE LDN')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = LIGHT

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run('Shop Sales Dashboard')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = BLACK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Redesign Planning Report')
r.font.size = Pt(16)
r.font.color.rgb = MID

# Gold rule
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(20)
pPr  = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot  = OxmlElement('w:bottom')
bot.set(qn('w:val'),   'single')
bot.set(qn('w:sz'),    '12')
bot.set(qn('w:space'), '0')
bot.set(qn('w:color'), 'E8C547')
pBdr.append(bot)
pPr.append(pBdr)

meta = [
    ('Date',     'June 2026'),
    ('Version',  '1.0 — Draft for review'),
    ('Project',  'shop-sales · github.com/ONE-LDN/shop-sales'),
    ('Prepared for', 'ONE LDN Management'),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rb = p.add_run(label + ':  ')
    rb.bold = True
    rb.font.size = Pt(10)
    rb.font.color.rgb = MID
    rv = p.add_run(val)
    rv.font.size = Pt(10)
    rv.font.color.rgb = LIGHT

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════
add_heading('1. Executive Summary', level=1, space_before=0)

add_body(
    'ONE LDN operates a café and merchandise shop at its gym facility. Sales are processed '
    'through WodBoard and tracked manually via CSV upload into a custom Supabase-backed dashboard. '
    'This report documents the planning and rationale behind a structural redesign of that '
    'dashboard — from a five-tab layout that had grown organically over time into a purposeful '
    'four-tab system built around four operational questions: what sold, what we have, what needs '
    'ordering, and how we are performing financially.'
)

add_heading('Objective', level=3, space_before=12, space_after=4)
add_callout(
    'To redesign the shop sales dashboard so that each section answers a single, clearly-defined '
    'question for a decision maker — eliminating overlap, surfacing the most operationally '
    'relevant signals on the default view, and providing a foundation for physical stock '
    'reconciliation that the business has needed from the start.',
    color_hex='F7F7F0', left_color='E8C547'
)

add_body(
    'The redesign was informed by a review of the metrics used by best-in-class inventory '
    'management systems, mapped against the specific operational context of a single-site gym shop '
    'with two distinct product types: high-frequency café consumables on a fortnightly order cycle, '
    'and longer-lead-time merchandise on an eight-week cycle.'
)

add_heading('Key Outcomes', level=3, space_before=12, space_after=4)
bullets = [
    ('Tab count reduced from 5 to 4. ', 'Overlap between the old Overview and Top Sellers tabs is eliminated.'),
    ('A default weekly report view. ', 'The landing screen now answers "what sold this week?" with week-on-week comparisons, replacing a static all-time summary.'),
    ('Stock reconciliation introduced. ', 'A new tab connects physical counts to the theoretical stock chain — a database table (shop_stock_takes) for this purpose already existed but had never been connected to the UI.'),
    ('Finance consolidated. ', 'Revenue trend charts and gross profit data previously split across two tabs are merged into a single monthly-review view.'),
    ('Health visibility on the primary screen. ', 'Stock urgency counts (Order Now / Order Soon / Healthy / Excess) and days-of-stock-remaining per product are visible on the default tab, removing the need to navigate away for a stock health check.'),
]
for bold, rest in bullets:
    add_bullet(rest, bold_prefix=bold)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 2. BACKGROUND
# ════════════════════════════════════════════════════════════════════
add_heading('2. Background', level=1)

add_heading('The Business Context', level=2, space_before=10, space_after=6)
add_body(
    'ONE LDN\'s shop covers two distinct trading areas: a café operation selling coffees, '
    'protein shakes, energy drinks, and snack bars; and a merchandise range including branded '
    'apparel and accessories. The two categories have meaningfully different inventory dynamics — '
    'café products turn over weekly and are ordered on a rolling two-week cycle from suppliers '
    'such as Muscle Finesse, while merchandise carries longer lead times and is ordered on an '
    'eight-week cycle from suppliers including Tapstitch.'
)
add_body(
    'Sales are processed through WodBoard and exported weekly as a CSV file. Deliveries are '
    'logged either manually or by uploading supplier PDF invoices, which the system parses '
    'automatically. All data is stored in Supabase.'
)

add_heading('Why a Redesign Was Needed', level=2, space_before=10, space_after=6)
add_body(
    'The original dashboard was built incrementally, with each new requirement added as an '
    'additional tab or element. Over time this produced five tabs with overlapping content, '
    'unclear primary actions, and a default view that did not reflect the most common reason '
    'for opening the dashboard — understanding the current week\'s trading and stock position.'
)
add_body('Specific structural problems identified in the previous design:')
problems = [
    'Overview and Top Sellers both answered variations of "what is selling?" in different formats, requiring users to visit two tabs for the same question.',
    'Stock History served two audiences simultaneously: operational stock reconciliation (delivery vs sold) and financial margin review (GP%), which belong in different contexts.',
    'The default Overview tab displayed metrics (all-time top products, revenue chart, category donut) that were useful for orientation but not for daily or weekly operational decisions.',
    'No physical stock count workflow existed in the UI, despite a purpose-built database table (shop_stock_takes) already being available in Supabase.',
    'Stock health status — days of stock remaining, urgency flags — was only visible inside the Order Prediction tab, requiring navigation away from the primary view to assess the situation.',
]
for p in problems:
    add_bullet(p)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 3. RESEARCH AND BENCHMARK CRITERIA
# ════════════════════════════════════════════════════════════════════
add_heading('3. Research and Benchmark Criteria', level=1)

add_body(
    'Before planning the redesign, the most useful metrics in best-in-class inventory management '
    'systems were identified. These formed the evaluation criteria for both the existing dashboard '
    'and the proposed new structure.'
)

add_heading('Industry Benchmark Metrics', level=2, space_before=10, space_after=6)
simple_table(
    ['Metric', 'Priority', 'Notes'],
    [
        ('Top-selling products and categories',                 'Core',    'By units and revenue, with period comparison'),
        ('Sell-through rate',                                   'Core',    'Units sold ÷ (delivered + opening stock)'),
        ('Days of inventory remaining',                         'Core',    'Stock ÷ weekly velocity'),
        ('Stock health flags — at-risk, healthy, overstocked',  'Core',    'Three-tier classification'),
        ('Revenue by week and month',                           'Core',    'Trend view over time'),
        ('Fastest and slowest moving items',                    'High',    'Velocity-based ranking'),
        ('Reorder alerts based on stock and velocity',          'Core',    'Urgency-flagged reorder list'),
        ('Gross profit % by product',                           'High',    'Requires cost price data'),
        ('Physical stock count reconciliation',                 'High',    'Actual vs theoretical'),
    ],
    col_widths=[8.0, 2.5, 5.5]
)

add_heading('Practical First-Version Requirements', level=2, space_before=10, space_after=6)
add_body(
    'From the benchmark research, five requirements were defined for a practical first version '
    'of the redesigned dashboard:'
)
reqs = [
    ('Headline KPIs', 'Revenue, units sold, gross margin where available, stock on hand — broken down by product and variant.'),
    ('Top products table', 'Ranked by units sold, showing revenue, sell-through rate, and days of stock remaining.'),
    ('Stock health panel', 'A clear summary of how many products are low-stock, healthy, overstocked, or excess.'),
    ('Trend chart', 'Revenue and units sold over time, segmented by category.'),
    ('Reorder list', 'Products below threshold or projected to run out, with suggested order quantities.'),
]
for bold, rest in reqs:
    add_bullet(rest, bold_prefix=bold + ' — ')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 4. THE PREVIOUS DESIGN — 5-TAB STRUCTURE
# ════════════════════════════════════════════════════════════════════
add_heading('4. The Previous Design — Five-Tab Structure', level=1)

add_body(
    'The original dashboard comprised five tabs. The following summarises each tab\'s purpose '
    'and the structural issues that informed the decision to redesign.'
)

simple_table(
    ['Tab', 'Purpose', 'Key Issues'],
    [
        ('Overview',         'KPI cards, revenue bar chart, category donut, top-10 product table', 'Mixed all-time and current-period data; overlaps with Top Sellers; stock alerts buried here rather than in Order Prediction'),
        ('Top Sellers',      'Brand and collection hierarchy tables for Café and Merch separately', 'Duplicates the "what\'s selling" question already partially answered by Overview; brand hierarchy not available on the default view'),
        ('Trends',           'Stacked line charts showing weekly revenue and units by category', 'Useful content isolated on a dedicated tab; not integrated with financial review'),
        ('Order Prediction', 'Velocity-based reorder table with covers bar and urgency flags',  'Strong and well-executed; appropriately isolated; correct as-is'),
        ('Stock History',    'Monthly delivery reconciliation plus gross profit % by product',   'Two distinct audiences (operational and financial) served by one tab; no physical count entry; GP% data hard to find'),
    ],
    col_widths=[3.0, 5.5, 7.5]
)

add_heading('What Worked Well in the Previous Design', level=2, space_before=10, space_after=6)
strong = [
    'The Order Prediction engine — velocity calculation, covers bar, urgency flags, and tab badge — was well-conceived and ahead of most entry-level commercial tools.',
    'PDF invoice parsing for Muscle Finesse and Tapstitch deliveries, reducing manual data entry significantly.',
    'The café/merch cycle-time distinction (two-week vs eight-week order cycles) correctly modelled the different inventory dynamics.',
    'The product lookup table structure (category, brand, cost price, pack size, min stock) provided a solid master data foundation.',
    'The exclusion logic for in-house products (house-made shakes, coffee) correctly scoped what the system should and should not attempt to predict.',
]
for s in strong:
    add_bullet(s)

add_heading('Coverage Against Benchmark Metrics', level=2, space_before=10, space_after=6)
simple_table(
    ['Metric', 'Status', 'Comment'],
    [
        ('Top-selling products and categories', ('Present', False, GREEN),   'Present but not as a week-on-week view'),
        ('Sell-through rate',                   ('Missing', True,  RED),     'Not calculated anywhere in the system'),
        ('Days of inventory remaining',          ('Present', False, GREEN),  'Order Prediction covers bar — strong'),
        ('Stock health flags',                   ('Partial', False, AMBER),  'No overstocked/excess flag; no health summary panel'),
        ('Revenue by week and month',            ('Partial', False, AMBER),  'Weekly on Trends tab; no daily; monthly via period selector'),
        ('Fastest/slowest movers',               ('Partial', False, AMBER),  'Velocity badges in Top Sellers; not on primary view'),
        ('Reorder alerts',                       ('Present', False, GREEN),  'Order Prediction — well executed'),
        ('Gross profit %',                       ('Partial', False, AMBER),  'Present in Stock History; data-dependent on cost entry'),
        ('Physical stock reconciliation',        ('Missing', True,  RED),    'Database table existed; no UI'),
    ],
    col_widths=[5.5, 2.5, 8.0]
)
add_callout(
    'Overall coverage against benchmark criteria: approximately 60–65%. The velocity and reorder '
    'engine was the standout component. The main gaps were on the reporting side — sell-through '
    'rate, overstocked flagging, and physical stock count integration.',
    color_hex='FFF8F0', left_color='E8C547'
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 5. THE REDESIGN — PROPOSED FOUR-TAB STRUCTURE
# ════════════════════════════════════════════════════════════════════
add_heading('5. The Redesign — Proposed Four-Tab Structure', level=1)

add_heading('Organising Principle', level=2, space_before=10, space_after=6)
add_body(
    'The redesign is structured around four sequential questions a manager asks when reviewing '
    'the business. Each tab answers exactly one question, with no overlap between them.'
)
simple_table(
    ['#', 'Tab', 'Question Answered', 'Replaces'],
    [
        ('1', 'Weekly Report',        'What sold this week?',       'Overview + Top Sellers'),
        ('2', 'Stock Reconciliation', 'What stock do we have?',     'Stock History (operational rows)'),
        ('3', 'Order Prediction',     'What needs reordering?',     'Order Prediction (unchanged)'),
        ('4', 'Finance',              'How are we performing?',     'Trends + Stock History (GP data)'),
    ],
    col_widths=[0.8, 3.8, 4.8, 5.0]
)

add_heading('Tab 1 — Weekly Report', level=2, space_before=10, space_after=6)
add_body(
    'The default landing tab. Replaces the old Overview and absorbs the Top Sellers product '
    'hierarchy. The primary user is someone opening the dashboard at the start or end of the week '
    'to understand what traded and whether anything needs immediate attention.'
)
add_body('Key elements:')
wr_items = [
    'Week navigation arrows and a dropdown to move between weeks.',
    'Stock health summary strip: four count badges (Order Now / Order Soon / Healthy / Excess Stock), each linking directly to the Order Prediction tab pre-filtered to that status.',
    'Per-category KPI cards (Coffee, Shakes, Bars & Snacks, Functional Drinks, Merch) showing units sold, revenue, and week-on-week percentage change.',
    'Product table grouped by brand and expandable to product level, sorted by units sold or revenue. Includes a Days Left column (from the velocity calculation) and a Current Stock column.',
    'Category and segment filters (All Café, individual categories, All Merch).',
]
for item in wr_items:
    add_bullet(item)

add_heading('Tab 2 — Stock Reconciliation', level=2, space_before=10, space_after=6)
add_body(
    'A new tab with no direct equivalent in the previous design. This is the connection between '
    'the theoretical stock position (calculated from deliveries minus sales) and the physical '
    'reality on the shelf. It is intended for use on a monthly cadence rather than weekly.'
)
add_body('Key elements:')
sr_items = [
    'Week navigation to review any historical week.',
    'Per-product reconciliation table: Opening Balance → + Deliveries → − Sold → = Theoretical → Closing Count (click to enter) → Variance.',
    'Opening balance automatically chains from the previous week\'s closing count.',
    'Inline count entry: clicking a Closing Balance cell opens a popover with a number input and an optional notes field.',
    'Monthly variance summary below the week table: aggregates variances across all weeks in a month where counts were entered, showing total shrinkage or surplus per product.',
    'Café / Merch segment toggle. In-house products (coffee, house-made shakes) are excluded — consistent with the Order Prediction exclusion logic.',
]
for item in sr_items:
    add_bullet(item)
add_callout(
    'The shop_stock_takes table for storing physical counts already existed in the Supabase '
    'database before the redesign. It had zero rows — the infrastructure was in place but '
    'had never been connected to a UI. This tab completes that connection.',
    color_hex='F0F7F0', left_color='1A7A3C'
)

add_heading('Tab 3 — Order Prediction', level=2, space_before=10, space_after=6)
add_body(
    'Carried over unchanged from the previous design. This is the strongest component of the '
    'original dashboard and required no structural changes — only a correction to its position '
    'in the tab order (previously tab 2; correctly tab 3, after reconciliation).'
)
add_body('Retained features:')
op_items = [
    'Velocity-based covers calculation: average units per week over the last four weeks, expressed as days of stock remaining.',
    'Visual covers bar (colour-coded by urgency) per product row.',
    'Urgency flags: Order Now (café < 7 days, merch < 14 days), Order Soon, Sufficient.',
    'Suggested order quantity based on cycle length (2 weeks for café, 8 weeks for merch).',
    'Tab badge showing the live "Order Now" count — visible from any tab.',
    'Minimum stock unit override per product (from the product lookup table).',
    'Café / Merch segment toggle and status filter (All / Order Now / Order Soon / Sufficient).',
]
for item in op_items:
    add_bullet(item)

add_heading('Tab 4 — Finance', level=2, space_before=10, space_after=6)
add_body(
    'A consolidated monthly review tab, merging the old Trends chart with the gross profit data '
    'from Stock History. Intended for management review rather than daily operational use.'
)
add_body('Key elements:')
fi_items = [
    'Weekly revenue by category — stacked line chart showing all historical weeks, with per-category toggle.',
    'Gross profit table by product with month picker: Product, Category, Units Sold, Revenue, COGS, GP%.',
    'Month-total GP% in the table footer.',
    'Cost price editing accessible from this tab (as it was from Stock History).',
]
for item in fi_items:
    add_bullet(item)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 6. PROTOTYPE ITERATIONS
# ════════════════════════════════════════════════════════════════════
add_heading('6. Prototype Iterations', level=1)

add_body(
    'Two prototype files were produced during the planning phase. Neither is connected to '
    'the live Supabase database — both use static mock data for layout and interaction validation.'
)

add_heading('Iteration 1 — prototype.html', level=2, space_before=10, space_after=6)
add_body(
    'An early-stage layout exploration. Introduced the concept of a dark-theme dashboard '
    'matching the ONE LDN brand aesthetic, and established the core visual language: '
    'category colour coding (red for coffee, gold for shakes, green for functional drinks, '
    'blue for merch), KPI card format, and the covers bar visualisation from Order Prediction. '
    'This file served as the visual foundation but did not yet define the tab structure.'
)

add_heading('Iteration 2 — prototype-structure.html', level=2, space_before=10, space_after=6)
add_body(
    'The primary structure prototype. Built to validate the four-tab layout, week navigation, '
    'expandable brand table, stock reconciliation flow, and the Finance chart. Key interactions '
    'implemented in static JavaScript:'
)
iter2_items = [
    'Week-by-week navigation using prev/next arrows and a dropdown selector.',
    'Brand-expandable product table with sort controls (units / revenue) and category filter.',
    'Stock reconciliation table with a popover count-entry UX — clicking a closing balance cell opens an input with a notes field; saving triggers a flash confirmation.',
    'Opening balance chaining from the previous week\'s closing count.',
    'Finance tab: interactive line chart with per-category legend toggles, and a month-picker GP table.',
    'Order Prediction: covers bar, urgency flags, status filter, and Order button per row.',
]
for item in iter2_items:
    add_bullet(item)

add_heading('Audit Findings and Revisions', level=2, space_before=10, space_after=6)
add_body(
    'Iteration 2 was formally audited against the benchmark criteria. The audit identified '
    'four high-priority issues and four medium-priority issues.'
)
add_body('High-priority issues addressed in the revised prototype:')
hi_items = [
    ('Tab order transposed. ', 'The prototype had Order Prediction at position 2 and Stock Reconciliation at position 3, reversing the intended flow. Corrected to: Weekly Report → Stock Reconciliation → Order Prediction → Finance.'),
    ('Days Left column missing from Weekly Report. ', 'The covers calculation existed in Order Prediction but was not visible on the default tab. Added as a colour-coded pill column in the product table.'),
    ('No stock health summary on the primary view. ', 'Order Now / Order Soon counts required navigating to a separate tab. Added as a four-badge health strip above the KPI cards, with click-through links to Order Prediction filtered by status.'),
    ('Monthly variance summary absent from Stock Reconciliation. ', 'Variance was calculable per week but not aggregated by month. Added a summary table below the weekly reconciliation table, aggregating all weeks with entered counts.'),
]
for bold, rest in hi_items:
    add_bullet(rest, bold_prefix=bold)

add_body('Medium-priority issues identified but not yet implemented:')
med_items = [
    'Sell-through rate not calculated — requires units sold ÷ (delivered + opening stock).',
    'No overstocked / excess flag in Order Prediction for products with more than 28 days cover (café) or 84 days cover (merch).',
    'No total-week revenue headline above the category KPI cards.',
    'Merch variant grouping (size, colourway) not demonstrated — data exists in the live system.',
]
for item in med_items:
    add_bullet(item)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 7. DECISIONS AND RATIONALE
# ════════════════════════════════════════════════════════════════════
add_heading('7. Key Design Decisions', level=1)

decisions = [
    (
        'Weekly navigation rather than a rolling-period selector',
        'The original dashboard used a rolling-period control (4 / 8 / 12 weeks or a specific month). '
        'The redesign replaces this on the primary tab with a direct week-by-week navigator. '
        'The rationale: the most common operational question is about a specific week, not a rolling average. '
        'The Finance tab retains period selection for trend analysis.'
    ),
    (
        'Stock health as count badges, not a bar chart',
        'A bar chart was considered for the stock health panel. Count badges were chosen instead '
        'because stock health is categorical (three or four states), not continuous. Badges '
        'communicate the urgency directly and double as navigation shortcuts.'
    ),
    (
        'Monthly cadence for stock reconciliation',
        'A weekly stock count cadence was considered impractical. The design allows closing counts '
        'to be entered for any week at any time, with the monthly variance summary aggregating '
        'only weeks where counts were actually recorded. This accommodates infrequent physical '
        'counts without breaking the chain.'
    ),
    (
        'Order Prediction unchanged',
        'The reorder engine was identified as the strongest component of the original design and '
        'required no structural changes. Rewriting it would have introduced risk with no benefit.'
    ),
    (
        'Café and merch in separate segments, not separate tabs',
        'Earlier designs had separate tabs for café and merch. The redesign uses a Café / Merch '
        'toggle within each tab. The rationale: the manager reviewing stock health cares about '
        'the overall picture first; the segment split is a filter, not a destination.'
    ),
    (
        'Excluding in-house products from reconciliation',
        'Coffee, house-made shakes, and services are excluded from the Stock Reconciliation tab '
        'for the same reason they are excluded from Order Prediction: they are not stocked products. '
        'This keeps the tab focused on items that are actually ordered and received from suppliers.'
    ),
]
for title, body in decisions:
    add_heading(title, level=3, space_before=10, space_after=4)
    add_body(body)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 8. COMPARISON AGAINST COMMERCIAL TOOLS
# ════════════════════════════════════════════════════════════════════
add_heading('8. Comparison Against Commercial Tools', level=1)

add_body(
    'The redesigned dashboard was assessed against the capabilities of entry-level commercial '
    'inventory management tools (Lightspeed Retail, Square for Retail, Vend).'
)
simple_table(
    ['Capability', 'Commercial', 'This Dashboard'],
    [
        ('Velocity-based reorder prediction',        '✓', '✓'),
        ('Visual stock coverage indicator',          '✓', '✓'),
        ('Physical stock count workflow',            '✓', '✓ (new tab)'),
        ('Gross profit % by product',                '✓', '✓ (data-dependent)'),
        ('Revenue trend chart by category',          '✓', '✓'),
        ('Sell-through rate',                        '✓', '✗ (planned)'),
        ('Overstocked / excess flag',                '✓', '✗ (planned)'),
        ('Real-time stock deduction on sale',        '✓', '✗ (CSV lag — by design)'),
        ('Automated purchase order generation',      '✓', '✗'),
        ('Dual café / merch cycle-time logic',       '✗', '✓ (custom)'),
        ('PDF invoice parsing for deliveries',       '✗', '✓ (custom)'),
        ('WodBoard CSV sales import',                '✗', '✓ (custom)'),
    ],
    col_widths=[8.5, 2.5, 5.0]
)
add_body(
    'The custom dashboard matches or exceeds commercial tools on the features most relevant to '
    'this business: velocity-based reordering, dual cycle-time logic, and automated invoice '
    'parsing. It is behind on sell-through rate and overstocked detection, both of which are '
    'technically straightforward additions at the next iteration.'
)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 9. NEXT STEPS
# ════════════════════════════════════════════════════════════════════
add_heading('9. Next Steps', level=1)

add_heading('Immediate — Prototype to Production', level=2, space_before=10, space_after=6)
immediate = [
    'Implement the four-tab structure in the live index.html, migrating existing logic into the new tab functions.',
    'Connect the Stock Reconciliation tab to the live shop_stock_takes Supabase table.',
    'Populate cost prices in the product lookup table to unlock gross profit reporting.',
    'Validate the reconciliation opening-balance chain with real delivery and sales data.',
]
for item in immediate:
    add_bullet(item)

add_heading('Short-Term — Remaining Gaps', level=2, space_before=10, space_after=6)
short = [
    'Add sell-through rate as a column in the Weekly Report product table.',
    'Add an overstocked / excess flag to the Order Prediction table for products above 28-day (café) or 84-day (merch) cover.',
    'Add a total-week revenue headline above the category KPI cards on the Weekly Report.',
    'Define merch variant grouping (size / colourway) for the Weekly Report product table.',
]
for item in short:
    add_bullet(item)

add_heading('Longer Term', level=2, space_before=10, space_after=6)
longer = [
    'Gross margin headline KPI on the Finance tab.',
    'Slow-movers callout: products with zero sales in the current week but positive estimated stock.',
    'Performance optimisation: parallelise the Supabase data loading calls on initialisation (currently sequential, making 9+ round-trips).',
    'Customer data utilisation: customer name and email are captured in shop_sales but unused.',
]
for item in longer:
    add_bullet(item)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
pPr  = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top  = OxmlElement('w:top')
top.set(qn('w:val'),   'single')
top.set(qn('w:sz'),    '4')
top.set(qn('w:space'), '4')
top.set(qn('w:color'), 'CCCCCC')
pBdr.append(top)
pPr.append(pBdr)
r = p.add_run('ONE LDN Shop Dashboard — Redesign Planning Report   ·   June 2026   ·   Internal document')
r.font.size = Pt(9)
r.font.color.rgb = LIGHT

# ── Save ─────────────────────────────────────────────────────────
doc.save('/home/user/shop-sales/shop-dashboard-redesign-report.docx')
print("Done: shop-dashboard-redesign-report.docx")
